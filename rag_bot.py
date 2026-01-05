import streamlit as st
import os
import json
import io
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
import threading
from docx import Document as DocxDocument  # 워드 파일 읽는 도구

# --- 1. 클라우드 환경 설정 ---
if "GOOGLE_API_KEY" in st.secrets:
    os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]
    SLACK_BOT_TOKEN = st.secrets["SLACK_BOT_TOKEN"]
    SLACK_APP_TOKEN = st.secrets["SLACK_APP_TOKEN"]
    GCP_KEY_DICT = json.loads(st.secrets["gcp_service_account"]["json_key"], strict=False)
else:
    st.error("🚨 Secrets 설정이 필요합니다!")
    st.stop()

# --- 2. RAG 두뇌 클래스 (워드 파일 지원 버전) ---
class CompanyBrain:
    def __init__(self):
        self.vector_store = None
        self.total_docs_count = 0
        # Gemini 3.0 Preview (안 되면 1.5-flash로 변경)
        self.llm = ChatGoogleGenerativeAI(model="gemini-3-flash-preview", temperature=0.3)
        self.load_db()

    def load_db(self):
        """구글 시트, 구글 문서, 그리고 MS Word(.docx)까지 모두 읽어옵니다."""
        print("📥 그룹디 지식 DB 동기화 중 (시트+문서+워드)...")
        
        # ▼▼▼ 폴더 ID 유지 ▼▼▼
        TARGET_FOLDER_ID = "1_sddYuhDRy1plDrCyA8GtKItQqVj4ULf" 
        
        try:
            scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
            creds = ServiceAccountCredentials.from_json_keyfile_dict(GCP_KEY_DICT, scope)
            
            client = gspread.authorize(creds)
            drive_service = build('drive', 'v3', credentials=creds)
            
            # 검색 쿼리: 시트 OR 구글문서 OR 워드파일
            query = f"'{TARGET_FOLDER_ID}' in parents and (mimeType = 'application/vnd.google-apps.spreadsheet' or mimeType = 'application/vnd.google-apps.document' or mimeType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document') and trashed = false"
            
            results = drive_service.files().list(q=query, fields="files(id, name, mimeType)").execute()
            items = results.get('files', [])

            if not items:
                print("⚠️ 폴더 안에 파일이 없습니다.")
                return

            documents = []
            
            for item in items:
                file_id = item['id']
                file_name = item['name']
                mime_type = item['mimeType']
                
                try:
                    # 1. 스프레드시트 (Google Sheets)
                    if 'spreadsheet' in mime_type:
                        sh = client.open_by_key(file_id) 
                        for worksheet in sh.worksheets():
                            title = worksheet.title
                            records = worksheet.get_all_records()
                            for row in records:
                                content_str = f"[시트: {file_name}-{title}] " + " / ".join([f"{k}: {v}" for k, v in row.items()])
                                documents.append(Document(page_content=content_str))
                                
                    # 2. 구글 문서 (Google Docs)
                    elif 'application/vnd.google-apps.document' in mime_type:
                        content = drive_service.files().export(fileId=file_id, mimeType='text/plain').execute().decode('utf-8')
                        if len(content.strip()) > 10:
                            doc_str = f"[문서: {file_name}] \n{content}"
                            documents.append(Document(page_content=doc_str))

                    # 3. 워드 파일 (.docx) - ⭐ 새로 추가된 기능
                    elif 'wordprocessingml.document' in mime_type:
                        # 워드 파일은 다운로드해서 메모리에서 엽니다
                        request = drive_service.files().get_media(fileId=file_id)
                        file_content = io.BytesIO()
                        downloader = MediaIoBaseDownload(file_content, request)
                        done = False
                        while done is False:
                            status, done = downloader.next_chunk()
                        
                        file_content.seek(0)
                        doc = DocxDocument(file_content)
                        full_text = [para.text for para in doc.paragraphs]
                        content = '\n'.join(full_text)
                        
                        if len(content.strip()) > 10:
                            doc_str = f"[Word파일: {file_name}] \n{content}"
                            documents.append(Document(page_content=doc_str))
                            print(f"📄 Word 파일 읽기 성공: {file_name}")
                            
                except Exception as e:
                    print(f"⚠️ '{file_name}' 읽기 실패: {e}")
                    continue

            if documents:
                self.total_docs_count = len(documents)
                embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
                self.vector_store = FAISS.from_documents(documents, embeddings)
                print(f"✅ 총 {len(documents)}개의 지식을 학습했습니다.")
            else:
                print("⚠️ 읽어올 데이터가 없습니다.")

        except Exception as e:
            print(f"❌ DB 로딩 실패: {e}")

    def ask(self, query):
        if not self.vector_store:
            return "아직 지식 DB가 준비되지 않았어요.", []
            
        prompt_template = f"""
        당신은 '그룹디(GroupD)'의 유능하고 센스 있는 AI 비서입니다.
        
        [현재 DB 현황]:
        - 학습된 지식 데이터 총 개수: {self.total_docs_count}건
        (사용자가 "전체 몇 개야?"라고 물으면 위 숫자를 답하세요.)

        [행동 지침]:
        1. 질문이 [회사의 지식]에 있는 업무 내용이라면, 정확하고 전문적으로 답변하세요.
        2. 질문이 일상 대화라면 친절하고 재치 있게 대화하세요.
        3. 답변은 항상 '해요체'(존댓말)로 정중하고 친절하게 하세요.

        [회사의 지식]:
        {{context}}

        [사용자 질문]:
        {{question}}

        [답변]:
        """
        PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

        qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            retriever=self.vector_store.as_retriever(search_kwargs={"k": 4}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        result = qa_chain.invoke({"query": query})
        return result["result"], result["source_documents"]

# 전역 캐시
@st.cache_resource
def get_brain():
    return CompanyBrain()

brain = get_brain()

# --- 3. 슬랙 봇 로직 ---
app = App(token=SLACK_BOT_TOKEN)

@app.message(".*")
def handle_message(message, say):
    query = message['text']
    try:
        answer, sources = brain.ask(query)
        source_text = ""
        if sources: 
            source_text = "\n\n📚 *참고 문서:*"
            for i, doc in enumerate(sources):
                preview = doc.page_content[:60].replace("\n", " ")
                source_text += f"\n> {i+1}. {preview}..."
            
        say(text=f"{answer}{source_text}", thread_ts=message['ts'])
    except Exception as e:
        say(f"❌ 에러 발생: {e}", thread_ts=message['ts'])

# --- 4. 메인 실행 ---
def run_slack_bot():
    try:
        handler = SocketModeHandler(app, SLACK_APP_TOKEN)
        handler.start()
    except Exception as e:
        print(f"봇 실행 에러: {e}")

st.title("🤖 그룹디(GroupD) 지식 봇 컨트롤러")
st.info("이 화면이 켜져 있으면 봇이 작동합니다.")

if st.button("🔄 지식 DB 업데이트"):
    brain.load_db() 
    st.success("최신 데이터를 반영했습니다!")

if 'bot_thread' not in st.session_state:
    bot_thread = threading.Thread(target=run_slack_bot, daemon=True)
    bot_thread.start()
    st.session_state.bot_thread = bot_thread
